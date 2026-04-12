# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUT_DIR = Path("results") / "thesis_draft"
DOCX_PATH = OUT_DIR / "毕业论文摘要与关键词.docx"
TXT_PATH = OUT_DIR / "毕业论文摘要与关键词.txt"
ALT_DOCX_PATH = OUT_DIR / "毕业论文摘要与关键词_修订版.docx"
ALT_TXT_PATH = OUT_DIR / "毕业论文摘要与关键词_修订版.txt"


CN_ABSTRACT = (
    "随着物流配送网络规模不断扩大以及客户服务时效要求不断提高，多配送中心带时间窗车辆路径问题逐渐成为物流优化研究中的重要内容。"
    "该问题同时涉及客户分配、车辆路径构造、容量约束和时间窗约束，具有较强的组合复杂性和实际工程意义。"
    "本文以多配送中心带时间窗车辆路径问题为研究对象，围绕分解法与整体法两条技术路线开展研究。"
    "在分解法框架下，将问题划分为客户分配和车场内路径优化两个阶段。第一阶段分别采用精确规划、广义指派、扫描修复以及基于蜣螂优化思想的修复分配方法进行求解与比较；"
    "第二阶段以蚁群算法为核心，针对带时间窗和异构车队条件下的路径优化问题，从多因子启发函数、信息素更新、自适应参数、群体分工、多蚁群协同以及蜣螂优化思想融合等方面对算法进行了改进。"
    "此外，本文还设计了整体法蚁群算法和整体法蜣螂优化算法，并与分解法进行对比分析。"
    "实验结果表明，在第一阶段中，精确规划方法具有较好的分配效果，而基于蜣螂优化思想的修复分配方法在后续路径求解衔接方面更具适用性。"
    "在第二阶段中，分解法整体优于整体法；在多种改进蚁群算法中，采用多因子启发信息、分层搜索机制、精英信息素强化和参数自适应优化的改进算法表现最好。"
    "综合实验结果可以看出，对于本文所研究的算例，采用“客户分配+车场内路径优化”的分解求解框架更为合适，"
    "将多因子启发、精英强化和群智能融合机制引入蚁群算法，能够有效提升多配送中心带时间窗车辆路径问题的求解效果。"
)

CN_KEYWORDS = "多配送中心车辆路径问题；时间窗；分解法；蚁群算法；蜣螂优化算法"

EN_ABSTRACT = (
    "With the expansion of logistics distribution networks and the increasing demand for service punctuality, the multi-depot vehicle routing problem with time windows has become an important topic in logistics optimization. "
    "This problem involves customer assignment, route construction, vehicle capacity constraints, and time window constraints, and therefore has strong combinatorial complexity and practical significance. "
    "This thesis studies the multi-depot vehicle routing problem with time windows from both decomposition-based and whole-network perspectives. "
    "Under the decomposition framework, the problem is divided into customer assignment and intra-depot route optimization. "
    "In the first stage, exact planning, generalized assignment, sweep-repair, and a repair assignment method inspired by dung beetle optimization are implemented and compared. "
    "In the second stage, ant colony optimization is improved through multi-factor heuristic information, pheromone update strategies, adaptive parameters, population role division, multi-colony cooperation, and the integration of dung beetle optimization ideas. "
    "Whole-network ant colony and dung beetle optimization methods are also designed for comparison. "
    "Experimental results show that the exact model performs well in customer assignment, while the dung-beetle-inspired repair assignment method is more suitable for the decomposition framework in subsequent routing optimization. "
    "For route optimization, the decomposition framework outperforms the whole-network framework, and the improved ant colony algorithm with multi-factor heuristic guidance, layered search, elite pheromone reinforcement, and adaptive parameter optimization achieves the best performance. "
    "The results indicate that the decomposition strategy is more appropriate for the current dataset, and that introducing multi-factor heuristics, elite reinforcement, and swarm-intelligence hybrid mechanisms can effectively improve the solution quality of the problem."
)

EN_KEYWORDS = "multi-depot vehicle routing problem; time windows; decomposition method; ant colony optimization; dung beetle optimization"


def set_run_font(run, font_name="SongTi", size=12, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.bold = bold


def add_title(doc, text, font_name, size):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, font_name, size, True)
    return p


def add_body(doc, text, font_name="宋体"):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, font_name, 12, False)
    return p


def add_keywords(doc, lead, words, font_name="宋体"):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(lead)
    set_run_font(r1, font_name, 12, True)
    r2 = p.add_run(words)
    set_run_font(r2, font_name, 12, False)
    return p


def build_doc():
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.18)
    sec.right_margin = Cm(3.18)

    add_title(doc, "摘  要", "黑体", 18)
    add_body(doc, CN_ABSTRACT, "宋体")
    add_keywords(doc, "关键词：", CN_KEYWORDS, "宋体")

    doc.add_page_break()

    add_title(doc, "ABSTRACT", "Times New Roman", 16)
    add_body(doc, EN_ABSTRACT, "Times New Roman")
    add_keywords(doc, "Key words: ", EN_KEYWORDS, "Times New Roman")

    saved_docx = DOCX_PATH
    try:
        doc.save(DOCX_PATH)
    except PermissionError:
        saved_docx = ALT_DOCX_PATH
        doc.save(saved_docx)

    text = "\n".join(
        [
            "摘  要",
            "",
            CN_ABSTRACT,
            "",
            f"关键词：{CN_KEYWORDS}",
            "",
            "ABSTRACT",
            "",
            EN_ABSTRACT,
            "",
            f"Key words: {EN_KEYWORDS}",
            "",
        ]
    )
    saved_txt = TXT_PATH
    try:
        TXT_PATH.write_text(text, encoding="utf-8")
    except PermissionError:
        saved_txt = ALT_TXT_PATH
        saved_txt.write_text(text, encoding="utf-8")

    print(f"saved: {saved_docx}")
    print(f"saved: {saved_txt}")


if __name__ == "__main__":
    build_doc()

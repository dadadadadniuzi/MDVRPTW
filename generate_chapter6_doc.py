# -*- coding: utf-8 -*-
from pathlib import Path
import re

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE = Path("results") / "thesis_draft"
TXT_OUT = BASE / "第六章_实验结果与性能分析.txt"
DOCX_OUT = BASE / "第六章_实验结果与性能分析_规范排版版.docx"

DATA_FILE = Path("data") / "mydata.xlsx"
PHASE1 = Path("results") / "comparison" / "phase1_comparison.csv"
PHASE2 = Path("results") / "comparison" / "phase2_comparison.csv"
REPRO_135 = Path("results") / "repro_compare" / "improve1_3_5_repeats_summary.csv"
REPRO_62 = Path("results") / "repro_compare" / "improve6_2_repeats_summary.csv"

FIG_IMPROVE5 = [
    Path("results") / "improved5_aco" / "improved5_route_depot_1.png",
    Path("results") / "improved5_aco" / "improved5_route_depot_2.png",
    Path("results") / "improved5_aco" / "improved5_route_depot_3.png",
]
FIG_WHOLE_ACO_CONV = Path("results") / "whole_aco" / "whole_aco_convergence.png"


def set_run_font(run, name, size_pt, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size_pt)
    run.bold = bold


def set_paragraph_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_page_field(paragraph):
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    paragraph._p.append(fld_begin)
    paragraph._p.append(instr)
    paragraph._p.append(fld_end)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = "w:" + edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn("w:" + key), str(value))


def format_paragraph(paragraph, first_line=False):
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(23)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)


def add_text_paragraph(doc, text, kind="body"):
    p = doc.add_paragraph()
    format_paragraph(p, first_line=(kind == "body"))
    if kind == "chapter":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, "黑体", 18, bold=True)
    elif kind == "section":
        run = p.add_run(text)
        set_run_font(run, "黑体", 14, bold=True)
    elif kind == "subsection":
        run = p.add_run(text)
        set_run_font(run, "黑体", 12, bold=True)
    else:
        run = p.add_run(text)
        set_run_font(run, "宋体", 12)
    return p


def add_table_title(doc, text):
    p = doc.add_paragraph()
    format_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, "宋体", 10.5)


def add_figure_title(doc, text):
    p = doc.add_paragraph()
    format_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, "宋体", 10.5)


def add_data_table(doc, headers, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for j, header in enumerate(headers):
        cell = table.cell(0, j)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        format_paragraph(p)
        run = p.add_run(str(header))
        set_run_font(run, "宋体", 10.5, bold=True)

    for i, row in enumerate(rows, start=1):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            format_paragraph(p)
            run = p.add_run(str(value))
            set_run_font(run, "宋体", 10.5)

    last_row = len(table.rows) - 1
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            border = {
                "top": {"val": "single", "sz": "8" if r_idx == 0 else "4", "color": "000000"},
                "bottom": {"val": "single", "sz": "8" if r_idx == last_row else "4", "color": "000000"},
                "left": {"val": "single", "sz": "4", "color": "000000"},
                "right": {"val": "single", "sz": "4", "color": "000000"},
            }
            set_cell_border(cell, **border)
    return table


def add_picture_with_caption(doc, image_path, caption, width_cm=12.2):
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    format_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    add_figure_title(doc, caption)


def load_setting_tables():
    customers = pd.read_excel(DATA_FILE, sheet_name="Customers")
    depots = pd.read_excel(DATA_FILE, sheet_name="Depots")
    vehicles = pd.read_excel(DATA_FILE, sheet_name="Vehicles")
    tw = pd.read_excel(DATA_FILE, sheet_name="TimeWindows")

    depot_tw = tw[tw["category"] == "depot"].copy()
    depot_table = depots.merge(depot_tw, on="id", how="left")
    depot_table["坐标"] = depot_table.apply(lambda r: f"({int(r['x'])}, {int(r['y'])})", axis=1)
    depot_table["营业时间窗"] = depot_table.apply(lambda r: f"{r['start_time']:.1f} - {r['end_time']:.1f}", axis=1)
    depot_table = depot_table[["id", "坐标", "capacity", "营业时间窗"]]
    depot_table.columns = ["车场编号", "坐标", "容量上限", "营业时间窗"]

    vehicles_table = vehicles.copy()
    vehicles_table.columns = ["所属车场", "车型编号", "数量", "容量", "速度", "固定成本", "单位里程成本"]

    customer_tw = tw[tw["category"] == "customer"].copy()
    customers_table = customers.merge(customer_tw, on="id", how="left")
    customers_table["坐标"] = customers_table.apply(lambda r: f"({int(r['x'])}, {int(r['y'])})", axis=1)
    customers_table["时间窗"] = customers_table.apply(
        lambda r: f"{r['start_time']:.1f} - {r['end_time']:.1f}", axis=1
    )
    customers_table["服务时间"] = customers_table["service_time"].map(lambda x: f"{x:.1f}")
    customers_table = customers_table[["id", "坐标", "demand", "时间窗", "服务时间"]]
    customers_table.columns = ["客户编号", "坐标", "需求量", "时间窗", "服务时间"]

    core_params = pd.DataFrame(
        [
            ["数据来源", "data/mydata.xlsx"],
            ["车场数量", "3"],
            ["客户数量", "30"],
            ["蚂蚁数量 m", "60"],
            ["蚁群最大迭代次数", "400"],
            ["信息素因子 α", "1.0"],
            ["启发因子 β", "2.0"],
            ["信息素挥发率 ρ", "0.1"],
            ["信息素强度 Q", "100"],
            ["DBO 种群规模", "50"],
            ["DBO 最大迭代次数", "200"],
            ["时间窗惩罚系数", "100"],
            ["复现实验次数", "5"],
        ],
        columns=["参数名称", "取值"],
    )
    return depot_table, vehicles_table, customers_table, core_params


def load_result_tables():
    phase1 = pd.read_csv(PHASE1)
    phase2 = pd.read_csv(PHASE2)
    repro135 = pd.read_csv(REPRO_135)
    repro62 = pd.read_csv(REPRO_62)
    repro = pd.concat([repro135, repro62], ignore_index=True)

    table_phase1 = phase1[
        ["Method", "Total_Assignment_Distance", "Depot_Loads", "Total_Overload", "Time_Feasible_Customers"]
    ].copy()
    table_phase1.columns = ["方法", "总分配距离", "车场负载", "总超载量", "时间可行客户数"]
    table_phase1["总分配距离"] = table_phase1["总分配距离"].map(lambda x: f"{x:.2f}")
    table_phase1["总超载量"] = table_phase1["总超载量"].map(lambda x: f"{x:.2f}")

    table_phase2 = phase2[phase2["Type"] == "Decomposition"][
        ["Method", "Route_Count", "Total_Distance", "Total_Cost"]
    ].copy()
    table_phase2.columns = ["方法", "车辆数", "总里程", "总成本"]
    table_phase2["总里程"] = table_phase2["总里程"].map(lambda x: f"{x:.2f}")
    table_phase2["总成本"] = table_phase2["总成本"].map(lambda x: f"{x:.2f}")

    table_repro = repro[["Method", "Mean_Cost", "Var_Cost", "Mean_Distance"]].copy()
    table_repro.columns = ["方法", "平均成本", "成本方差", "平均里程"]
    table_repro["平均成本"] = table_repro["平均成本"].map(lambda x: f"{x:.3f}")
    table_repro["成本方差"] = table_repro["成本方差"].map(lambda x: f"{x:.3f}")
    table_repro["平均里程"] = table_repro["平均里程"].map(lambda x: f"{x:.3f}")

    table_whole = phase2[phase2["Type"] == "Whole-network"][
        ["Method", "Route_Count", "Total_Distance", "Total_Cost", "Avg_Load_Rate"]
    ].copy()
    table_whole.columns = ["方法", "车辆数", "总里程", "总成本", "平均装载率"]
    table_whole["总里程"] = table_whole["总里程"].map(lambda x: f"{x:.2f}")
    table_whole["总成本"] = table_whole["总成本"].map(lambda x: f"{x:.2f}")
    table_whole["平均装载率"] = table_whole["平均装载率"].map(lambda x: f"{x:.4f}")
    return table_phase1, table_phase2, table_repro, table_whole


def build_text():
    return """6 实验结果与性能分析

6.1 实验环境与评价指标

6.1.1 算例与基础数据设置
为验证本文所提出分解法与整体法求解框架的有效性，并分析不同客户分配策略和路径优化算法在多配送中心带时间窗车辆路径问题中的适用性，本文在统一数据集和统一程序环境下开展对比实验。实验算例来源于项目中的自定义数据文件，包含3个配送中心、30个客户点以及异构车辆信息。为了使后续实验结果具有清晰的数据背景，本节首先给出配送中心设置和客户、车辆整体设置。
从配送中心配置看，本文算例共设置3个车场，不同车场在空间位置、容量上限以及营业时间窗方面存在差异。车场1的容量上限为90，车场2的容量上限为130，车场3的容量上限为120；三个车场的营业时间窗分别为8.0-21.0、6.0-18.0和0.0-24.0。这样的设置使不同车场在可服务能力与时间弹性方面表现出差异，也使客户归属与路径构造问题更加贴近实际配送场景。
从客户与车辆整体设置看，本文算例共包含30个客户点，客户需求量分布在一定范围内，且客户时间窗存在明显差异，同时设置了统一的服务时间。车辆方面，不同配送中心配置了不同类型和数量的车辆，车辆在容量、速度和单位里程成本上并不完全相同，因此形成了异构车队条件。上述算例特征决定了算法不仅需要优化客户访问顺序，还必须同时兼顾客户归属、车型匹配以及时间窗可行性判断。

6.1.2 实验程序核心参数
相比学长论文的实验分析，本研究原始版本在实验程序参数说明方面仍显不足。为了增强实验过程的可复现性，本节进一步给出本文核心求解程序的统一参数设置。对于整体法与分解法中涉及的蚁群算法，本文统一采用60只蚂蚁和400次最大迭代次数；对于用于分配阶段和部分改进算法中的蜣螂优化部分，统一采用50个个体和200次最大迭代次数。基础参数中，信息素启发因子α取1.0，启发函数因子β取2.0，信息素挥发率ρ取0.1，信息素强度Q取100。为了控制迟到客户的惩罚力度，时间窗惩罚系数设置为100。
上述参数设置一方面参考了基础蚁群算法和蜣螂优化算法的常用经验范围，另一方面也结合了本文算例规模进行了稳定性调整。对于客户规模为30、车场数量为3的中小规模MDVRPTW算例而言，400次蚁群迭代已经能够使基础算法和大部分改进算法完成较为充分的搜索，而200次蜣螂迭代能够保证分配阶段和参数寻优阶段具有一定的全局探索能力。后续所有单次实验与复现实验，均在上述统一参数框架下进行，以保证不同算法之间具有可比性。

6.1.3 评价指标与统计口径
在实验评价指标方面，本文同时记录第一阶段客户分配结果和第二阶段路径优化结果。第一阶段主要考察总分配距离、各车场负载分布、总超载量以及时间可行客户数，用于评价不同分配策略在距离紧凑性、容量平衡性和可服务性方面的表现。第二阶段则重点考察车辆使用数量、总行驶里程、总成本以及多次重复实验下的均值和方差，用于分析不同路径优化算法的解质量和稳定性。
需要说明的是，分解法第二阶段在算法内部求解时还会考虑时间窗迟到惩罚，而整体法早期版本曾采用“固定出车成本+里程成本”的总成本口径。为了保证整体法和分解法结果表之间具有一致的比较基础，本文在结果统计阶段统一采用“路径距离×车辆单位里程成本”的方式重新计算整体法总成本，不再单独将固定出车成本计入最终对比表。因此，本章给出的整体法与分解法总成本结果，已经在统计口径上做了统一处理。

6.2 第一阶段客户-车场分配结果分析

6.2.1 不同分配方法的结果比较
第一阶段实验主要比较扫描修复方法、GAP启发式方法、精确数学规划方法以及DBO+Repair方法在客户-车场分配问题中的表现。表6.5给出了四种方法在总分配距离、车场负载、总超载量和时间可行客户数方面的对比结果。
从总分配距离指标看，Exact MILP取得了588.70的最优值，说明在“仅以分配距离最小”为目标的数学规划意义下，该方法能够提供理论最优边界。DBO+Repair的总分配距离为594.90，仅比Exact MILP高6.20；GAP方法为598.19；Sweep+Repair方法为618.59，相对较高。从数值上看，DBO+Repair已经明显优于几何基线分配和普通启发式分配。
从可行性角度看，四种方法在当前算例下均未出现超载现象，且30个客户全部通过时间可行性检查。这说明在本文设置的数据规模和修复机制下，四种分配方法均能够生成可行方案。但是，如果仅依据分配距离进行评价，仍不足以完全解释不同分配结构对后续路径优化的影响。对本文而言，第一阶段更重要的作用是为第二阶段构造结构清晰、局部紧凑、便于独立求解的客户簇。

6.2.2 DBO+Repair方法分析
结合第一阶段整体结果可以看出，DBO+Repair方法在分配距离、容量平衡和后续可实现性之间取得了较好的折中。该方法并不是单纯追求分配距离最短，而是在DBO动态权重搜索基础上，通过修复阶段主动规避可能影响车场内路径构造的结构性问题。相较于Exact MILP直接追求理论边界，DBO+Repair更强调面向两阶段求解框架的衔接效果。
这也是本文后续第二阶段统一采用DBO+Repair分配结果作为输入的主要原因。因为在分解法框架下，第一阶段的意义不仅在于“把客户分出去”，更在于“把客户合理地分出去”。如果分配结构不能兼顾后续时间窗和容量约束，即使第一阶段在距离上更优，也未必能在第二阶段获得更好的系统总成本。因此，DBO+Repair在本文中既是一个分配算法，也是连接分解法两阶段的关键桥梁。

6.3 第二阶段路径优化结果分析

6.3.1 基础算法与改进算法单次结果比较
第二阶段实验主要围绕车场内独立VRPTW路径优化展开。本文以基础ACO VRPTW为起点，逐步设计并实现Improve 1、Improve 2、Improve 3、Improve 4、Improve 5、Improve 6、Improve 6.2和Improve 6.3等改进算法，并与整体法结果进行统一比较。表6.6给出了分解法第二阶段各算法在车辆数、总里程和总成本上的单次实验结果。
从分解法内部对比看，基础ACO VRPTW的总成本为468.10；Improve 1将总成本降低至463.71，说明将DBO等级制度思想引入蚁群角色分工后，能够在一定程度上改善搜索质量；Improve 3的总成本为464.46，也优于基础算法，但提升幅度略低于Improve 1；Improve 4由于外层调参收益没有有效转化为路径质量提升，总成本上升至490.49，说明仅增加调参层次并不能保证带来更优结果。
在所有改进算法中，Improve 5的表现最好，其总里程为450.95，总成本为461.67，优于前述各版本。Improve 6.2虽然在机制上进一步引入了多蚁群协同和DBO融合控制，但总成本仍为461.79，未能超过Improve 5；Improve 6.3的总成本为470.43，也未能进一步取得提升。由此可见，在本文算例下，算法结构并不是越复杂越好，关键在于改进机制是否真正贴合VRPTW构解过程本身，而Improve 5在复杂度与求解效果之间取得了最好的平衡。
从机制层面看，Improve 5同时引入了多因子启发函数、异质蚂蚁分工、精英信息素强化以及DBO外层调参机制，形成了较为完整的协同结构。与Improve 1和Improve 3这类局部增强型改进相比，Improve 5不仅在状态转移阶段更重视时间窗紧迫度和容量匹配性，在信息素更新和参数适配层面也更加系统，因此最终表现出更强的单次求解能力。

6.3.2 改进算法复现实验与稳定性分析
为了避免仅凭单次实验结果判断算法优劣，本文进一步对代表性改进算法进行了5次复现实验，并统计平均成本、成本方差和平均里程。表6.7给出了Improve 1、Improve 3、Improve 5和Improve 6.2的重复实验统计结果。
从平均成本看，Improve 5的均值为465.834，低于Improve 3的466.169、Improve 1的468.340和Improve 6.2的470.195，说明Improve 5并非依赖偶然性的单次最优解，而是在多次重复实验意义下仍然保持较强的综合表现。换言之，Improve 5既能在单次实验中取得最优结果，也能在重复实验中保持较低的平均成本，说明其算法设计更具有稳健性。
从方差指标看，Improve 6.2的成本方差为8.423，小于Improve 5的20.928，表明多蚁群协同与内外层控制机制确实提高了搜索稳定性。但与此同时，Improve 6.2的平均成本仍高于Improve 5，这说明其稳定性提升尚未转化为更优的总体成本。相较之下，Improve 3的成本方差达到115.095，波动明显更大，表明仅依靠局部参数或信息素机制改进，仍难以在复杂约束场景下长期保持稳定输出。综合单次最优值、均值和方差三个维度，Improve 5更适合作为本文分解法第二阶段的主推算法。

6.3.3 路线可视化与收敛过程分析
为了使实验分析不只停留在表格对比层面，本文进一步给出最优分解法算法的路径可视化结果，并结合整体法基础算法的收敛曲线进行说明。图6.1至图6.3展示了Improve 5在三个车场内得到的配送路线图。从图中可以看出，经过第一阶段客户分配后，各车场负责的客户在空间上已经形成了较为清晰的局部服务区域，第二阶段路径优化主要在局部紧凑区域内完成线路组织。这说明分解法前后两阶段之间具有较强的配合关系，第一阶段不仅降低了第二阶段搜索空间，也为后续路线构造提供了较好的几何基础。
从车场内路径分布看，Improve 5得到的路线总体上没有出现明显的大范围折返和无效绕行，说明多因子启发机制与精英信息素强化能够较好地抑制不合理跳转。对于带时间窗和异构车队的VRPTW问题而言，路径图的紧凑性并不只意味着空间距离更短，也意味着客户服务次序与车辆能力之间形成了更自然的匹配关系。
图6.4给出了Whole ACO的收敛曲线。可以看到，整体法基础蚁群算法在前期迭代中成本下降较快，说明整体法在搜索初期具有一定的全局可行解构造能力；随着迭代继续进行，曲线后期逐渐趋于平缓，表明算法在当前参数设置下能够完成一定程度的解精化，但进一步下降空间已经有限。结合表6.8中的结果可以发现，整体法即使在基础Whole ACO版本下也仍未超过分解法中的Improve 5。这说明在本文研究对象下，先通过第一阶段完成客户-车场分配，再在局部子问题上实施高质量路径优化，更有利于稳定获得较低的系统总成本。

6.3.4 整体法与分解法结果对比
在统一总成本统计口径后，整体法与分解法之间的差异表现得更为清晰。Whole ACO、Whole DBO和Whole ACO Improve1的结果如表6.8所示。整体法各版本虽然能够完成全局求解，但从总成本水平看，均未超过分解法中的Improve 5，其中Whole DBO和Whole ACO Improve1与最优分解法之间的差距更为明显。
这一现象说明，整体法虽然理论上能够同步处理客户归属与路径构造，但其搜索空间显著大于分解法，因此对构解机制和参数控制的要求也更高。分解法通过第一阶段先完成客户分配，相当于提前缩小了第二阶段搜索空间，使改进蚁群算法能够更集中地针对车场内VRPTW进行优化。对于本文当前算例而言，这种“先分配、后优化”的框架更有利于稳定获得高质量解。
这说明，整体法虽然理论上能够同步处理客户归属与路径构造，但其搜索空间更大、约束耦合更强，因此对构解机制和参数控制的要求也更高。相比之下，分解法先通过第一阶段缩小问题规模，再在第二阶段针对车场内VRPTW进行深入优化，更容易形成稳定且高质量的求解过程。因此，在本文当前算例和算法体系下，分解法明显优于整体法，而Improve 5则是分解法路径优化阶段中表现最优的算法。

6.4 本章小结

本章围绕实验结果与性能分析，对本文提出的分解法与整体法求解框架进行了统一验证。与原始版本相比，本章补充了算例基础设置、实验程序核心参数、统一成本统计口径、路径可视化分析以及收敛过程分析，使实验章节的结构更加完整，也更接近学长论文中“参数—结果—图示—结论”相结合的分析方式。
实验结果表明，在第一阶段中，Exact MILP给出了总分配距离的理论最优边界，而DBO+Repair方法在距离、容量平衡与后续可实现性之间取得了更适合作为两阶段输入的折中方案。在第二阶段中，Improve 5在单次实验和重复实验中均表现出较强的综合优势，是本文分解法框架下当前效果最好的路径优化算法。整体法方面，各版本结果均未超过Improve 5，说明在本文当前研究对象和实验条件下，分解法仍然是更有优势的主要求解路线。"""


def build_doc():
    BASE.mkdir(parents=True, exist_ok=True)
    depot_table, vehicles_table, customers_table, core_params_table = load_setting_tables()
    phase1_table, phase2_table, repro_table, whole_table = load_result_tables()
    text = build_text()
    TXT_OUT.write_text(text, encoding="utf-8")

    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.tab_stops.add_tab_stop(Cm(15.0), WD_TAB_ALIGNMENT.RIGHT)
    hrun = hp.add_run("东北大学本科生毕业设计（论文）\t6 实验结果与性能分析")
    set_run_font(hrun, "宋体", 10)
    set_paragraph_border(hp)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(fp)

    table_inserted = {
        "6.1.1 算例与基础数据设置": [
            ("表6.1 配送中心设置", depot_table),
            ("表6.2 车辆配置表", vehicles_table),
            ("表6.3 客户与时间窗数据表", customers_table),
        ],
        "6.1.2 实验程序核心参数": [
            ("表6.4 实验程序核心参数", core_params_table),
        ],
        "6.2.1 不同分配方法的结果比较": [
            ("表6.5 第一阶段客户分配结果比较", phase1_table),
        ],
        "6.3.1 基础算法与改进算法单次结果比较": [
            ("表6.6 分解法第二阶段单次实验结果比较", phase2_table),
        ],
        "6.3.2 改进算法复现实验与稳定性分析": [
            ("表6.7 改进算法复现实验统计结果", repro_table),
        ],
        "6.3.4 整体法与分解法结果对比": [
            ("表6.8 整体法实验结果比较", whole_table),
        ],
    }

    figure_inserted = {
        "6.3.3 路线可视化与收敛过程分析": [
            ("图6.1 Improve 5 ACO车场1路径图", FIG_IMPROVE5[0], 11.8),
            ("图6.2 Improve 5 ACO车场2路径图", FIG_IMPROVE5[1], 11.8),
            ("图6.3 Improve 5 ACO车场3路径图", FIG_IMPROVE5[2], 11.8),
            ("图6.4 Whole ACO收敛曲线", FIG_WHOLE_ACO_CONV, 12.5),
        ],
    }

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if re.match(r"^\d+\s", line):
            add_text_paragraph(doc, line, "chapter")
        elif re.match(r"^\d+\.\d+\.\d+\s", line):
            add_text_paragraph(doc, line, "subsection")
            if line in table_inserted:
                for title, df in table_inserted[line]:
                    add_table_title(doc, title)
                    add_data_table(doc, list(df.columns), df.values.tolist())
            if line in figure_inserted:
                for caption, image_path, width in figure_inserted[line]:
                    add_picture_with_caption(doc, image_path, caption, width_cm=width)
        elif re.match(r"^\d+\.\d+\s", line):
            add_text_paragraph(doc, line, "section")
        else:
            add_text_paragraph(doc, line, "body")

    doc.save(DOCX_OUT)
    print(f"saved: {TXT_OUT}")
    print(f"saved: {DOCX_OUT}")


if __name__ == "__main__":
    build_doc()

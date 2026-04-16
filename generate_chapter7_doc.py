# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE = Path("results") / "thesis_draft"
TXT_OUT = BASE / "第七章_结论与展望.txt"
DOCX_OUT = BASE / "第七章_结论与展望_规范排版版.docx"
ALT_DOCX_OUT = BASE / "第七章_结论与展望_规范排版版_修订版.docx"


CHAPTER_TEXT = """7 结论与展望

本文围绕考虑时间窗的多配送中心车辆路径优化问题开展了较为系统的研究工作。针对该问题同时具有客户分配决策与车辆路径决策、并受到容量约束、时间窗约束和多车场结构共同影响的特点，本文分别从分解法与整体法两条技术路线出发，构建了相应的求解框架，并围绕蚁群优化算法开展了多种改进尝试。通过模型分析、算法设计与实验对比，较为完整地考察了不同求解思路在本文算例条件下的适用性。

本文的主要研究工作包括：

（1）围绕多配送中心带时间窗车辆路径问题建立了与本文研究对象相对应的模型描述，明确了客户需求、车辆容量、配送中心营业时间和客户时间窗等基本约束，并从分解法与整体法两个角度梳理了问题求解思路，为后续算法设计提供了统一的问题基础。

（2）在分解法框架下，将问题划分为客户—车场分配与车场内路径优化两个阶段。针对第一阶段，分别实现了扫描法、广义指派启发式、MILP 理论边界方法以及基于蜣螂优化思想的动态权重修复分配策略，并对不同分配方法的结果进行了比较分析。实验结果表明，单纯依赖几何邻近或理论最优分配边界并不一定能够为后续路径求解带来最优衔接效果，而结合动态权重与可行性修复的分配方法在本文算例中表现出更好的综合适配性。

（3）针对分解法第二阶段的车场内独立 VRPTW 路径优化问题，本文以蚁群优化算法为核心，围绕多因子启发函数、信息素更新机制、参数自适应调节、群体角色分工、多蚁群协同以及蜣螂优化思想融合等方面进行了多轮改进，先后构建了多个改进版本。实验比较表明，在本文已实现的算法中，基于 DBO 在线调参驱动的超启发式改进蚁群算法在总成本、总里程以及算法稳定性方面表现最好，是当前分解法框架下的最优方案。

（4）在整体法方面，本文分别设计并实现了整体蜣螂优化算法、整体蚁群算法以及融合等级制度思想的整体改进蚁群算法，对客户归属与路径构造进行联合求解。通过与分解法结果的统一对比可以看出，整体法虽然能够从全局视角直接处理客户分配与路径规划的耦合关系，但在当前算例规模和约束条件下，其搜索空间更大、可行解构造难度更高，因此总体结果仍弱于分解法。

（5）本文在统一实验环境下对不同分配方法、不同改进算法以及分解法与整体法的结果进行了系统比较，并对部分核心算法进行了重复实验统计。实验结果说明，针对本文所研究的算例，先进行客户分配、再进行车场内路径优化的分解框架更适合问题求解；同时，将多因子启发信息、精英强化机制和蜣螂优化思想合理融入蚁群算法，能够有效提升算法在复杂约束条件下的求解效果。

综合全文研究结果，可以得到以下主要结论：

第一，在多配送中心带时间窗车辆路径问题中，客户分配结果会直接影响后续路径优化的可行性与求解质量，因此第一阶段并不能简单视为几何聚类问题，而应更多考虑其与第二阶段路径构造之间的衔接关系。

第二，在本文算例条件下，分解法整体优于整体法。其主要原因在于，分解法通过先完成客户归属划分，有效缩小了第二阶段的搜索空间，使路径优化算法能够更集中地处理车场内部的时间窗与容量约束；而整体法虽然具有全局协调潜力，但在当前算法能力和算例规模下尚未充分释放这种优势。

第三，蚁群优化算法仍然是求解该类路径构造问题的有效方法，但基础蚁群算法在复杂约束条件下面临明显局限。通过引入多因子启发机制、精英信息素强化、自适应参数调节和群体分工思想，能够显著改善算法的搜索效率与解质量。

第四，将蜣螂优化思想用于蚁群算法改进，而不是简单替代蚁群算法，是本文实验中更有效的技术路径。尤其是在参数调优、角色分工和搜索多样性维护方面，DBO 思想为 ACO 的进一步改进提供了较好的支持。

针对本文已经完成的研究工作，结合当前结果与仍存在的不足，后续还可以从以下几个方面继续展开：

（1）进一步扩大算例规模，并增加不同客户分布模式、不同时间窗强度和不同车辆配置条件下的实验分析，以考察当前算法在更复杂场景中的稳定性与可扩展性。

（2）在模型层面进一步贴近实际配送环境，例如考虑动态订单插入、道路拥堵导致的时间依赖行驶时间、冷链运输约束以及更丰富的车辆调度规则，使模型更具工程应用价值。

（3）继续围绕改进蚁群算法开展更有针对性的机制优化，特别是在整体法框架下，加强客户归属决策与路径构造决策之间的信息联动，以进一步提升整体法的求解能力。

（4）尝试将深度学习、强化学习或更高效的超启发式方法与现有群智能算法结合，构建更具自适应能力的求解框架，并在统一实验平台下与本文已有算法进行系统比较。

总体而言，本文完成了对多配送中心带时间窗车辆路径问题的建模、算法设计与实验验证工作，并在分解法框架下得到了较为理想的求解结果。受时间和研究条件限制，相关工作仍有进一步完善的空间，但本文的研究结果能够为后续深入开展该类问题的优化研究提供一定参考。"""


def set_run_font(run, name, size_pt, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size_pt)
    run.bold = bold


def set_paragraph_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_page_field(paragraph):
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    paragraph._p.append(fld_begin)
    paragraph._p.append(instr)
    paragraph._p.append(fld_end)


def next_available_path(path: Path) -> Path:
    if not path.exists():
        return path
    stem = path.stem
    suffix = path.suffix
    idx = 2
    while True:
        candidate = path.parent / f"{stem}_{idx}{suffix}"
        if not candidate.exists():
            return candidate
        idx += 1


def format_body_paragraph(p):
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(23)
    p.paragraph_format.first_line_indent = Cm(0.74)


def build_doc():
    BASE.mkdir(parents=True, exist_ok=True)
    TXT_OUT.write_text(CHAPTER_TEXT, encoding="utf-8")

    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.tab_stops.add_tab_stop(Cm(15.0), WD_TAB_ALIGNMENT.RIGHT)
    hrun = hp.add_run("东北大学本科生毕业设计（论文）\t7 结论与展望")
    set_run_font(hrun, "宋体", 10)
    set_paragraph_border(hp)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    frun = fp.add_run()
    set_run_font(frun, "Times New Roman", 10)
    add_page_field(fp)

    for idx, line in enumerate(CHAPTER_TEXT.splitlines()):
        if not line.strip():
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(23)
        if idx == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            set_run_font(run, "黑体", 18, bold=True)
        else:
            p.paragraph_format.first_line_indent = Cm(0.74)
            run = p.add_run(line)
            set_run_font(run, "宋体", 12)

    try:
        doc.save(DOCX_OUT)
        saved = DOCX_OUT
    except PermissionError:
        alt = next_available_path(ALT_DOCX_OUT)
        doc.save(alt)
        saved = alt

    print(f"saved: {TXT_OUT}")
    print(f"saved: {saved}")


if __name__ == "__main__":
    build_doc()

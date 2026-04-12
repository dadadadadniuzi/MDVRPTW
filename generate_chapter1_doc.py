# -*- coding: utf-8 -*-
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


SRC = Path("results") / "thesis_draft" / "第一章_绪论.txt"
OUT = Path("results") / "thesis_draft" / "第一章_绪论_规范排版版.docx"


def set_run_font(run, name, size_pt, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size_pt)
    run.bold = bold


def set_paragraph_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_page_field(paragraph):
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    paragraph._p.append(fld_begin)
    paragraph._p.append(instr)
    paragraph._p.append(fld_end)


def classify(line: str):
    if not line.strip():
        return "blank"
    if line.startswith("第一章"):
        return "chapter"
    if re.match(r"^\d+\.\d+\.\d+", line):
        return "subsection"
    if re.match(r"^\d+\.\d+", line):
        return "section"
    return "body"


def build_doc():
    text = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()

    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), WD_TAB_ALIGNMENT.RIGHT)
    hrun = hp.add_run("东北大学本科生毕业设计（论文）\t1 绪论")
    set_run_font(hrun, "宋体", 10)
    set_paragraph_border(hp)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    frun = fp.add_run()
    set_run_font(frun, "Times New Roman", 10)
    add_page_field(fp)

    for line in text:
        kind = classify(line)
        if kind == "blank":
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.5

        if kind == "chapter":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            set_run_font(run, "黑体", 18, bold=True)  # 小二近似
        elif kind == "section":
            run = p.add_run(line)
            set_run_font(run, "黑体", 14, bold=True)  # 四号近似
        elif kind == "subsection":
            run = p.add_run(line)
            set_run_font(run, "黑体", 12, bold=True)  # 小四
        else:
            p.paragraph_format.first_line_indent = Cm(0.74)
            run = p.add_run(line)
            set_run_font(run, "宋体", 12)

    doc.save(OUT)
    print(f"saved: {OUT}")


if __name__ == "__main__":
    build_doc()

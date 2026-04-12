# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Cm


OUT_DIR = Path("results") / "reports"
DOCX_PATH = OUT_DIR / "彭骏杰第六次汇报3-17.docx"
TXT_PATH = OUT_DIR / "彭骏杰第六次汇报3-17.txt"


TITLE = "第六次汇报（3月17日）"

SECTIONS = [
    (
        "这一阶段主要完成的工作",
        [
            "从第五次汇报到现在，我主要还是继续做分解法第二阶段，重点是继续找比现在更好的蚁群算法。",
            "这一阶段我补做并整理了 Improve4、Improve5、Improve6 和 Improve6.2，另外也把整体法、复现实验、对比表、中期报告和论文第一章一起推进了。"
        ],
    ),
    (
        "Improve4",
        [
            "Improve4 主要是在前面分层蚂蚁和精英信息素更新的基础上，再把 DBO 用到参数整定上。",
            "它调的参数比较多，包括 alpha、beta、挥发率上下界、信息素上下界、变异概率、时间窗惩罚和精英增强系数。",
            "我原本希望它能通过自动调参把算法整体性能再往上推，但结果并不好，总里程 476.04，总成本 490.49。",
            "我目前的判断是，Improve4 的问题主要是调参维度有点多，参数之间互相耦合，DBO 虽然能找到一组可用参数，但不一定能稳定找到真正适合当前算例的组合，所以反而没有带来收益。"
        ],
    ),
    (
        "Improve5",
        [
            "Improve5 是我这一阶段最满意的一个版本。它的核心思路是：保留蚁群算法做路径构造，再借鉴蜣螂算法的分工思想，并用 DBO 去调关键参数。",
            "它主要有三点改动。第一，多因子启发函数，把距离、时间窗紧迫度和容量匹配度一起放进状态转移概率。第二，蚂蚁分层，不同蚂蚁承担强化、常规搜索、启发式搜索和扰动搜索。第三，信息素更新时更强调当前优解和全局优解，而不是所有蚂蚁平均贡献。",
            "这个版本的结果目前最好，总里程 450.95，总成本 461.67。",
            "我后面又做了 5 次复现，它的平均成本 465.834，方差 20.928。目前看，Improve5 仍然是分解法第二阶段最适合作为主打方案的算法。"
        ],
    ),
    (
        "Improve6",
        [
            "Improve6 主要是尝试 MACS，也就是多蚁群系统。我把蚁群拆成两个群体，一个偏向总成本，一个偏向车辆使用和平衡，然后分别维护信息素，再做一定的信息交换。",
            "这个版本想解决的问题是：单一蚁群容易搜索方向太单一，多蚁群可能能把搜索面拉开。",
            "结果上，Improve6 的总里程 455.06，总成本 467.46，没有超过 Improve5。",
            "我现在的理解是，MACS 这个方向是有研究价值的，但单独做成 Improve6 时，群体之间的协同收益还不够，反而削弱了对当前最优方向的集中强化。"
        ],
    ),
    (
        "Improve6.2",
        [
            "Improve6.2 是在 Improve6 的基础上继续往前走的一版。它保留了 MACS 双蚁群结构，同时把 DBO 更深地融入了进去。",
            "具体来说，这一版不只是外层用 DBO 调参数，还加了一层内层 DBO 控制器，在算法运行过程中动态调整两个蚁群的 alpha、beta 和信息素交换比例。",
            "单次结果上，Improve6.2 的总里程 451.08，总成本 461.79，已经非常接近 Improve5，只差 0.12。",
            "我又做了 5 次复现，它的平均成本 470.195，方差 8.423。这个结果说明它的波动比 Improve5 小，稳定性更好，但均值还没有压下来。所以我觉得 Improve6.2 现在更像一个很有潜力的框架，还需要继续简化和调优。"
        ],
    ),
    (
        "复现实验与结果判断",
        [
            "为了不只看单次结果，我把 Improved ACO、Improved3 ACO、Improve5 和 Improve6.2 都做了 5 次复现。",
            "目前看，Improve5 的均值最好，单次最好值也最好；Improve6.2 的优势是方差更小，结果更稳，但均值还没超过 Improve5。",
            "所以如果现在就要定论文第二阶段主打算法，我认为还是应该先用 Improve5；如果要继续做创新点，最值得往下挖的是 Improve6.2。"
        ],
    ),
    (
        "这一阶段完成的其他工作",
        [
            "第一，我把第一阶段和第二阶段的结果整理成了统一对比表，方便直接看各算法优劣。",
            "第二，我把整体法的蚁群算法、蜣螂算法和整体改进蚁群算法也都实现并做了对比，目前结果还是明显不如分解法。",
            "第三，我补做了中期报告，并开始写毕业论文第一章，目前第一章和规范排版版文档都已经整理出来了。"
        ],
    ),
    (
        "下一步计划",
        [
            "下一步我准备继续沿 Improve6.2 这个方向往下做，重点看能不能在保留框架创新性的同时，把均值继续往下压。",
            "另外，我也会开始把现在已经稳定下来的算法设计、对比实验和结果分析逐步写进论文正文，这样后面不会集中堆积写作任务。"
        ],
    ),
]


def set_run_font(run, font_name="宋体", size=12, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.bold = bold


def build_doc():
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.18)
    sec.right_margin = Cm(3.18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    set_run_font(r, "黑体", 16, True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("彭骏杰")
    set_run_font(r2, "宋体", 12, False)

    for heading, paras in SECTIONS:
        hp = doc.add_paragraph()
        hp.paragraph_format.space_before = Pt(6)
        hp.paragraph_format.space_after = Pt(3)
        hr = hp.add_run(heading)
        set_run_font(hr, "黑体", 14, True)

        for para in paras:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.74)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(para)
            set_run_font(r, "宋体", 12, False)

    doc.save(DOCX_PATH)

    parts = [TITLE, "彭骏杰", ""]
    for heading, paras in SECTIONS:
        parts.append(heading)
        parts.extend(paras)
        parts.append("")
    TXT_PATH.write_text("\n".join(parts), encoding="utf-8")

    print(f"saved: {DOCX_PATH}")
    print(f"saved: {TXT_PATH}")


if __name__ == "__main__":
    build_doc()

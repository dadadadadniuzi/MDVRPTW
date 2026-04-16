# -*- coding: utf-8 -*-
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


TXT_OUT = Path("results") / "thesis_draft" / "第五章_整体法求解框架设计与实现.txt"
DOCX_OUT = Path("results") / "thesis_draft" / "第五章_整体法求解框架设计与实现_规范排版版.docx"


def set_run_font(run, name, size_pt, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size_pt)
    run.bold = bold


def set_paragraph_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_page_field(paragraph):
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    paragraph._p.append(fld_begin)
    paragraph._p.append(instr)
    paragraph._p.append(fld_end)


def classify(line: str):
    if not line.strip():
        return "blank"
    if line.startswith("5 "):
        return "chapter"
    if re.match(r"^\d+\.\d+\.\d+", line):
        return "subsection"
    if re.match(r"^\d+\.\d+", line):
        return "section"
    return "body"


def clean_xml_text(text: str) -> str:
    return "".join(ch for ch in text if ch in "\t\n\r" or ord(ch) >= 32)


def build_text():
    return r"""5 整体法求解框架设计与实现

5.1 整体法核心思想与连续转离散编码机制

与第四章采用“先分配、后路径优化”的分解法不同，整体法不再预先固定客户归属，而是将客户分配、车辆选择和访问顺序统一纳入同一个求解框架。在该框架下，算法需要在全局范围内同时回答三个问题：客户应由哪个配送中心承担服务、应由哪一类车辆执行配送任务、以及各客户在车辆路径中的访问先后顺序如何安排。这样处理的优点在于，算法在搜索过程中能够直接利用全局信息，避免分阶段决策可能带来的局部最优问题；但与此同时，整体法的决策维度更高，搜索空间显著扩大，对编码方式、解码方式和约束处理机制提出了更高要求。

本文整体法求解框架的基本流程为：首先读取全部配送中心、客户、车辆和时间窗数据，并在系统层面建立客户距离矩阵、配送中心到客户的距离矩阵以及车辆池；随后采用统一编码表示客户的全局优先级信息，并通过解码机制将连续型编码转换为离散的客户服务顺序；然后在解码过程中动态完成客户归属判定、车辆分配与路径构造；最后依据总成本对候选解进行评价，并通过群智能算法不断迭代更新，直至达到终止条件。

由于整体法直接作用于多配送中心带时间窗车辆路径问题，而该问题本身属于典型的离散组合优化问题，因此本文没有直接让元启发式算法在离散解空间上进行硬编码搜索，而是采用“连续搜索、离散解码”的设计思想。具体而言，在整体法中，无论是 Whole DBO 还是 Whole ACO，其核心搜索对象都可以抽象为一个长度为客户数的连续向量。设客户总数为 $n$，则连续编码可表示为

$$X=(x_1,x_2,\dots,x_n) \qquad (5.1)$$

其中，$x_i$ 表示客户 $i$ 对应的连续优先级值。算法并不直接将该向量视为最终路径，而是根据各分量数值大小对客户进行排序，得到全局访问优先级序列

$$\pi=\mathrm{argsort}(X) \qquad (5.2)$$

式中，$\pi$ 为由小到大排序后得到的客户排列结果。排序后的序列并不等于最终解，而是作为解码器的输入。解码器依次读取客户序列中的元素，在当前已有路径和剩余车辆资源条件下，判断该客户更适合插入哪一条已有路径，或由哪一辆新车辆开启新路径。由此，连续编码便被转换为满足多配送中心结构的离散路径方案。

在整体法框架中，连续转离散编码机制的关键不在于排序本身，而在于排序后的动态解码规则。对任一待插入客户，算法并不预先指定其所属配送中心，而是同时考察所有可行路径插入方案与可行新建路径方案。若某客户插入已有路径后满足车辆容量、客户时间窗以及配送中心营业时间约束，则计算其新增距离成本；若插入不可行，则继续尝试以其他车辆从相应配送中心出发为其新建路径。于是，客户归属并不是在编码阶段单独决定的，而是在解码阶段伴随车辆调度和路径延伸过程共同确定。

这种连续转离散编码机制具有两方面作用。其一，它使整体法中的 DBO 和 ACO 能够共享较统一的解表示方式，从而便于比较不同算法的全局搜索能力。其二，它避免了直接对多配送中心路径结构进行复杂离散编码所带来的实现困难，使算法可以先在连续空间中完成搜索，再通过统一解码器映射到可行路径方案。需要指出的是，这种编码方式虽然提升了实现的统一性，但也意味着解的优劣高度依赖解码策略本身，尤其是在时间窗约束较强、车辆类型较多时，若解码规则不够强，就容易出现可行解质量不高的问题。

因此，整体法的本质可以理解为：以连续型优先级向量作为统一搜索载体，以动态解码器作为离散化桥梁，在全局范围内同时完成客户分配、车辆调度和路径构造。基于这一总体设计，本文分别实现了全局蜣螂优化算法和全局蚁群算法，并在后续小节中对两类方法的具体结构进行说明。

5.2 全局蜣螂优化算法(Whole DBO)设计

在整体法框架下，本文首先实现了全局蜣螂优化算法 Whole DBO，用于直接求解多配送中心带时间窗车辆路径问题。该算法延续了 DBO 在连续空间中的群体搜索思想，并结合前述连续转离散编码机制，将每个蜣螂个体表示为一个客户优先级连续向量。与第四章第一阶段 DBO 只优化配送中心权重不同，Whole DBO 直接面向全局路径方案进行搜索，因此其求解对象从“分配权重”转变为“全局客户访问优先级”。

Whole DBO 的初始种群由若干个随机向量构成，每个向量通过统一解码器映射为一组完整路径。设第 $p$ 个个体对应的解码结果为 $S_p$，则其适应度可记为

$$F_p=C(S_p) \qquad (5.3)$$

其中，$C(S_p)$ 表示该路径方案的总成本。在本文实现中，总成本主要由各条路径的距离成本构成，即车辆行驶距离与单位运输成本的乘积之和。若某个个体在解码阶段无法为全部客户构造可行路径，则将其适应度置为无穷大，从而在种群筛选阶段予以淘汰。

在解码过程中，Whole DBO 按照优先级序列逐个处理客户。对于当前客户，算法首先尝试将其插入已有路径末端，并检查插入后的容量约束、客户时间窗和配送中心营业时间约束是否成立；若插入可行，则计算新增距离代价，并在所有可行插入方案中选择代价最小的一项。若当前客户无法插入已有路径，则进一步尝试使用尚未启用的车辆从某一配送中心出发为其新建路径；若新建路径仍不可行，则当前个体被判定为无效。由此可见，Whole DBO 的解码并非简单排序后直接切段，而是一个带约束判断的逐步构造过程。

在种群更新方面，Whole DBO 保留了滚球、育雏、小蜣螂和偷窃四类个体的分工思想。滚球个体更强调向当前全局最优位置靠拢，用于增强开发能力；育雏个体围绕当前最优个体附近进行局部扰动搜索；小蜣螂个体通过更大的随机扰动扩大搜索范围；偷窃个体则借助其他个体附近的信息进行跳跃式更新，以维持种群多样性。若将当前最优个体记为 $X^{best}$，最差个体记为 $X^{worst}$，则滚球个体的更新可概括为

$$X_i^{t+1}=X_i^t+a_1r_1(X^{best}-X_i^t)-a_2r_2(X^{worst}-X_i^t) \qquad (5.4)$$

其中，$r_1$ 和 $r_2$ 为随机向量，$a_1$、$a_2$ 为调节系数。育雏与小蜣螂个体则更多依赖高斯扰动和随机同伴信息，以在最优区域附近进行更细致搜索。

Whole DBO 的优势在于它能够在全局层面直接处理客户归属与路径结构，不需要像分解法那样先人为固定车场边界；但其不足也较为明显。由于一个个体向量需要通过复杂解码器映射为完整路径，适应度评价成本较高，而且解码结果高度依赖车辆可用性和时间窗约束，因此在搜索过程中较容易出现可行解比例偏低、局部更新收益有限的问题。从本文实验结果看，Whole DBO 的整体表现明显弱于分解法下的主打算法，也弱于整体法中的基础 Whole ACO。这说明，对于当前算例而言，DBO 直接用于整体路径构造时，其全局搜索能力并未充分转化为高质量解。

尽管如此，Whole DBO 仍然具有研究意义。它为本文提供了一个不经过分阶段拆分、直接面向全局多配送中心问题的基准求解器，也使得后续 Whole ACO 与其改进算法能够在统一的整体法框架下进行对比。换言之，Whole DBO 在本文中的定位更偏向于整体法对照算法和全局搜索机制代表，而不是最终最优方案。

5.3 全局蚁群算法(Whole ACO)设计

为了进一步探索整体法在多配送中心带时间窗车辆路径问题中的适用性，本文在 Whole DBO 之外又设计了全局蚁群算法 Whole ACO。与 Whole DBO 通过连续优先级向量进行全局搜索不同，Whole ACO 更强调在构解阶段利用信息素和启发因子逐步形成路径方案。其核心思想是：不预先划分客户所属车场，而是由蚂蚁在全局范围内依次为客户选择服务车辆和服务次序，并在迭代中利用优质解不断强化有效边结构。

与分解法下的车场内蚁群算法相比，Whole ACO 的差异主要有两点。第一，搜索对象不再是固定客户簇内部的单车场路径，而是同时包含客户归属和跨配送中心车辆选择的全局方案。第二，启发信息不再只由客户间距离决定，还需要考虑客户对不同配送中心的初始接近程度以及车辆类型差异。正因为如此，Whole ACO 在整体法中不仅是一个路径算法，也承担了部分客户归属决策功能。

5.3.1 全局启发因子与多车型动态调度机制

Whole ACO 在结构上仍保留了蚁群算法“信息素引导 + 启发式构解”的基本模式，但在启发因子设计和车辆调度逻辑上进行了整体法适配。首先，在启发矩阵构建时，算法并不只计算客户之间的距离倒数，还额外考虑客户从最近配送中心出发的初始接近程度。对于客户 $j$，若其到最近配送中心的距离记为 $d_{0j}^{min}$，则起始启发值可写为

$$\eta_{0j}=\frac{1}{d_{0j}^{min}+\varepsilon} \qquad (5.5)$$

对任意客户 $i$ 和客户 $j$，客户间启发信息仍保留为

$$\eta_{ij}=\frac{1}{d_{ij}+\varepsilon} \qquad (5.6)$$

其中，$\varepsilon$ 为防止分母为零的极小正数。通过这种设计，算法在路径起始阶段能够优先关注更接近某一配送中心的客户，而在路径延伸阶段则继续利用客户间空间邻近关系。

在构解过程中，Whole ACO 并不预先固定某一辆车服务某一批客户，而是维护一个全局车辆池。每次需要开启新路径时，算法首先从尚未使用的车辆中筛选出能够服务至少一个未访问客户的可行车辆；随后根据车辆所属配送中心、容量、速度和单位成本进行评分，从中选出更适合当前搜索状态的车辆。设候选车辆为 $v$，其所属配送中心到最近可服务客户的距离为 $d_v^{near}$，单位运输成本为 $c_v$，固定出车成本为 $f_v$，则车辆选择评分可概括为

$$\mathrm{score}(v)=d_v^{near}\cdot c_v+\lambda f_v \qquad (5.7)$$

其中，$\lambda$ 为固定成本权重。在本文重新统一成本口径后的结果统计中，固定成本不再计入最终比较值，但在车辆选择阶段仍可作为辅助偏好信息。由此，Whole ACO 并不是随机启动车辆，而是通过动态评分机制优先启用更适合当前客户分布的车辆类型。

在路径延伸阶段，算法对每个未访问客户执行容量、时间窗与回仓可行性检查。若某客户由当前车辆服务后无法在配送中心关闭前返回，则该客户不会进入候选集合。对于可行客户，算法结合信息素、启发函数和时间窗紧迫度计算选择概率，并通过轮盘赌完成下一客户选择。设当前状态到客户 $j$ 的信息素为 $\tau_{ij}$，启发信息为 $\eta_{ij}$，时间紧迫度为 $u_j$，则其综合选择概率可写为

$$P_{ij}=\frac{\tau_{ij}^{\alpha}\eta_{ij}^{\beta}u_j}{\sum\limits_{h\in \Omega}\tau_{ih}^{\alpha}\eta_{ih}^{\beta}u_h} \qquad (5.8)$$

其中，$\Omega$ 为当前可行客户集合。由此，Whole ACO 在整体法框架下仍然保持了较强的构解能力，同时借助车辆池与配送中心距离信息完成全局动态调度。

5.3.2 融合等级制度的改进全局蚁群算法 (Whole ACO Improve)

为了进一步提升整体法下蚁群算法的搜索多样性，本文在 Whole ACO 基础上设计了融合等级制度的改进全局蚁群算法，即 Whole ACO Improve。该算法借鉴了前文 Improve 1 中的角色分工思想，但其作用场景从“车场内独立子问题”扩展到“全局多配送中心联合搜索”。其核心目标是在整体法中缓解单一蚁群容易早熟收敛的问题，使不同类型蚂蚁在全局路径构造阶段承担不同搜索任务。

具体而言，Whole ACO Improve 将蚂蚁划分为滚球组、育雏组、小蜣螂组和带扰动的探索组。滚球组更强调信息素作用，倾向于沿当前较优结构继续开发；育雏组使用标准参数完成主体搜索；小蜣螂组降低信息素权重、增强启发信息作用，偏向于快速构造局部较优路径；探索组则在一定概率下随机选择候选客户，用于增加解空间覆盖范围。若第 $k$ 类蚂蚁使用的参数记为 $(\alpha_k,\beta_k)$，则状态转移概率可写为

$$P_{ij}^{(k)}=\frac{\tau_{ij}^{\alpha_k}\eta_{ij}^{\beta_k}}{\sum\limits_{h\in \Omega}\tau_{ih}^{\alpha_k}\eta_{ih}^{\beta_k}} \qquad (5.9)$$

不同角色参数的差异，决定了各类蚂蚁在搜索过程中承担的功能并不相同。

在信息素更新方面，Whole ACO Improve 保留了代内优质解强化思路。与基础 Whole ACO 相比，该改进版本会对滚球组中的优质解给予更高权重，同时对当前代最优解再进行额外强化，以加快优质路径结构在种群中的传播。由此，Whole ACO Improve 希望在整体法场景中兼顾“优质结构快速强化”和“种群多样性维持”两方面目标。

但从本文实验结果看，Whole ACO Improve 并未取得比基础 Whole ACO 更好的整体表现。其原因主要在于：在整体法中，客户归属、车辆启用和路径顺序本就耦合紧密，搜索空间比车场内独立优化大得多。此时即使引入等级制度和异质参数，也未必能像分解法那样稳定发挥作用。对于当前算例而言，基础 Whole ACO 已经能够依赖较直接的全局启发信息构造相对紧凑的可行路径，而 Whole ACO Improve 增加的额外随机性和角色差异，反而在一定程度上削弱了整体搜索方向的集中性。这也说明，分解法中有效的改进机制在迁移到整体法时，并不一定能够获得同等收益。

5.4 本章小结

本章围绕整体法求解框架，对多配送中心带时间窗车辆路径问题的全局求解思路进行了系统说明。首先，给出了整体法的核心思想，并阐述了连续编码与离散解码相结合的统一表示机制，说明了客户归属、车辆选择和路径构造如何在同一框架下联合完成。随后，分别介绍了 Whole DBO、Whole ACO 及 Whole ACO Improve 的算法结构与实现逻辑，分析了它们在整体法场景下的编码方式、构解机制与改进重点。

通过本章分析可以看出，整体法的优势在于能够在全局范围内直接处理客户分配与路径优化之间的耦合关系，但其缺点也同样明显，即搜索空间大、解码复杂、可行性判断成本高。Whole DBO 强调全局连续搜索能力，Whole ACO 则更依赖启发式构解与信息素强化，而 Whole ACO Improve 则尝试通过等级制度增强搜索多样性。整体而言，这些方法为本文建立了较完整的整体法研究框架，并为后续实验章节中分解法与整体法的性能比较提供了算法基础。"""


def build_doc():
    text = build_text()
    TXT_OUT.parent.mkdir(parents=True, exist_ok=True)
    TXT_OUT.write_text(text, encoding="utf-8")

    lines = text.splitlines()
    doc = Document()

    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.tab_stops.add_tab_stop(Cm(15.0), WD_TAB_ALIGNMENT.RIGHT)
    hrun = hp.add_run("东北大学本科生毕业设计（论文）\t5 整体法求解框架设计与实现")
    set_run_font(hrun, "宋体", 10)
    set_paragraph_border(hp)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    frun = fp.add_run()
    set_run_font(frun, "Times New Roman", 10)
    add_page_field(fp)

    for line in lines:
        line = clean_xml_text(line)
        kind = classify(line)
        if kind == "blank":
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(23)

        if kind == "chapter":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            set_run_font(run, "黑体", 18, bold=True)
        elif kind == "section":
            run = p.add_run(line)
            set_run_font(run, "黑体", 14, bold=True)
        elif kind == "subsection":
            run = p.add_run(line)
            set_run_font(run, "黑体", 12, bold=True)
        else:
            p.paragraph_format.first_line_indent = Cm(0.74)
            run = p.add_run(line)
            set_run_font(run, "宋体", 12)

    doc.save(DOCX_OUT)
    print(f"saved: {TXT_OUT}")
    print(f"saved: {DOCX_OUT}")


if __name__ == "__main__":
    build_doc()

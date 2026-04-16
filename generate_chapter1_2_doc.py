# -*- coding: utf-8 -*-
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE = Path("results") / "thesis_draft"

CHAPTERS = [
    {
        "txt": BASE / "第一章_绪论.txt",
        "docx": BASE / "第一章_绪论_规范排版版.docx",
        "alt_docx": BASE / "第一章_绪论_规范排版版_修订版.docx",
        "header_right": "1 绪论",
        "chapter_no": "1",
    },
    {
        "txt": BASE / "第二章_相关理论与问题描述.txt",
        "docx": BASE / "第二章_相关理论与问题描述_规范排版版.docx",
        "alt_docx": BASE / "第二章_相关理论与问题描述_规范排版版_修订版.docx",
        "header_right": "2 相关理论与问题描述",
        "chapter_no": "2",
    },
]

MERGED_DOCX = BASE / "第一二章_规范排版版.docx"
MERGED_ALT_DOCX = BASE / "第一二章_规范排版版_修订版.docx"


def next_available_path(path: Path) -> Path:
    if not path.exists():
        return path
    stem = path.stem
    suffix = path.suffix
    parent = path.parent
    index = 2
    while True:
        candidate = parent / f"{stem}_{index}{suffix}"
        if not candidate.exists():
            return candidate
        index += 1


def set_run_font(run, name, size_pt, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size_pt)
    run.bold = bold


def set_paragraph_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_page_field(paragraph):
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    paragraph._p.append(fld_begin)
    paragraph._p.append(instr)
    paragraph._p.append(fld_end)


def configure_section(section, header_right):
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.tab_stops.add_tab_stop(Cm(15.0), WD_TAB_ALIGNMENT.RIGHT)
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after = Pt(0)
    text = f"东北大学本科生毕业设计（论文）\t{header_right}"
    run = hp.add_run(text)
    set_run_font(run, "宋体", 10)
    set_paragraph_border(hp)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    run = fp.add_run()
    set_run_font(run, "Times New Roman", 10)
    add_page_field(fp)


def configure_normal_style(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)


def classify(line: str, chapter_no: str):
    if not line.strip():
        return "blank"
    if re.match(rf"^{re.escape(chapter_no)}\s", line):
        return "chapter"
    if re.match(r"^\d+\.\d+\.\d+", line):
        return "subsection"
    if re.match(r"^\d+\.\d+", line):
        return "section"
    return "body"


def add_formatted_paragraph(doc, line: str, kind: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(23)

    if kind == "chapter":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        set_run_font(run, "黑体", 18, bold=True)
    elif kind == "section":
        run = p.add_run(line)
        set_run_font(run, "黑体", 14, bold=True)
    elif kind == "subsection":
        run = p.add_run(line)
        set_run_font(run, "黑体", 12, bold=True)
    else:
        p.paragraph_format.first_line_indent = Cm(0.74)
        run = p.add_run(line)
        set_run_font(run, "宋体", 12)


def build_single_doc(chapter):
    lines = chapter["txt"].read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_normal_style(doc)
    configure_section(doc.sections[0], chapter["header_right"])

    for line in lines:
        kind = classify(line, chapter["chapter_no"])
        if kind == "blank":
            continue
        add_formatted_paragraph(doc, line, kind)

    try:
        doc.save(chapter["docx"])
        return chapter["docx"]
    except PermissionError:
        alt = next_available_path(chapter["alt_docx"])
        doc.save(alt)
        return alt


def build_merged_doc():
    doc = Document()
    configure_normal_style(doc)
    first = True

    for chapter in CHAPTERS:
        lines = chapter["txt"].read_text(encoding="utf-8").splitlines()
        if first:
            section = doc.sections[0]
            first = False
        else:
            section = doc.add_section(WD_SECTION_START.NEW_PAGE)
        configure_section(section, chapter["header_right"])

        for line in lines:
            kind = classify(line, chapter["chapter_no"])
            if kind == "blank":
                continue
            add_formatted_paragraph(doc, line, kind)

    try:
        doc.save(MERGED_DOCX)
        return MERGED_DOCX
    except PermissionError:
        alt = next_available_path(MERGED_ALT_DOCX)
        doc.save(alt)
        return alt


def main():
    BASE.mkdir(parents=True, exist_ok=True)
    for chapter in CHAPTERS:
        saved = build_single_doc(chapter)
        print(f"saved: {saved}")
    merged_saved = build_merged_doc()
    print(f"saved: {merged_saved}")


if __name__ == "__main__":
    main()
